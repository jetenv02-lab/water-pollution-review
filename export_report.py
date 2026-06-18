# -*- coding: utf-8 -*-
"""審查報告匯出產生器 (v2 — 多選包含項版)。

對外 API:
    build_export(target, app_data, findings=None, options=None, base_name=None)
        target: "internal" / "vendor" / "json"
        options 支援的 keys (對應 11 個包含項):
            include_reminder        💡 提醒 findings
            include_system_error    ⚠️ 系統錯誤 findings (debug)
            include_unit_table      📋 全廠單元表
            include_water_quality   💧 各單元水質
            include_design_metrics  📐 設計參數體檢
            include_removal_rate    📉 削減率分項
            include_rule_match      📚 觸發規則對照
            include_rule_text       📜 規則條文節錄
            include_raw_json_dump   🔬 原始 JSON dump (內嵌)
            include_pdf_notes       📝 PDF 註解 (reviewer_notes)
            include_default_values  🧪 業別預設值
            include_discharge_std   ⚖️ 放流水標準
        回傳 (bytes_data, file_name, mime)

對象 → 格式:
    internal  → Excel (多分頁)
    vendor    → Word (改善建議書, 廠商可編輯回覆)
    json      → JSON (整合 + 選用區段)
"""
from __future__ import annotations

import io
import json
from datetime import datetime


# ─────────────────────────────────────────
# 各對象的預設「包含內容」
# ─────────────────────────────────────────

DEFAULT_OPTIONS = {
    "internal": {
        "include_reminder": False,
        "include_system_error": False,
        "include_unit_table": True,
        "include_water_quality": True,
        "include_design_metrics": True,
        "include_removal_rate": True,
        "include_rule_match": True,
        "include_rule_text": False,
        "include_raw_json_dump": False,
        "include_pdf_notes": True,
        "include_default_values": False,
        "include_discharge_std": False,
    },
    "vendor": {
        "include_reminder": False,
        "include_system_error": False,
        "include_unit_table": True,
        "include_water_quality": False,
        "include_design_metrics": False,
        "include_removal_rate": False,
        "include_rule_match": True,
        "include_rule_text": True,
        "include_raw_json_dump": False,
        "include_pdf_notes": False,
        "include_default_values": False,
        "include_discharge_std": True,
    },
    "json": {
        "include_reminder": True,
        "include_system_error": True,
        "include_unit_table": True,
        "include_water_quality": True,
        "include_design_metrics": True,
        "include_removal_rate": True,
        "include_rule_match": True,
        "include_rule_text": False,
        "include_raw_json_dump": True,
        "include_pdf_notes": True,
        "include_default_values": True,
        "include_discharge_std": True,
    },
}


# 顯示用標籤 (給 UI 用)
OPTION_LABELS = {
    "include_reminder":       ("📊 基本",   "💡 提醒 findings"),
    "include_system_error":   ("📊 基本",   "⚠️ 系統錯誤 (debug)"),
    "include_unit_table":     ("📋 單元",   "全廠單元表"),
    "include_water_quality":  ("📋 單元",   "各單元水質"),
    "include_design_metrics": ("📋 單元",   "設計參數體檢"),
    "include_removal_rate":   ("📋 單元",   "削減率分項"),
    "include_rule_match":     ("📚 規則",   "觸發規則對照"),
    "include_rule_text":      ("📚 規則",   "規則條文節錄"),
    "include_raw_json_dump":  ("🔬 進階",   "原始 JSON dump"),
    "include_pdf_notes":      ("🔬 進階",   "PDF 註解 (reviewer)"),
    "include_default_values": ("🔬 進階",   "業別預設值"),
    "include_discharge_std":  ("🔬 進階",   "放流水標準"),
}


def get_default_options(target: str) -> dict:
    """取得指定對象的預設 options。"""
    return dict(DEFAULT_OPTIONS.get(target, DEFAULT_OPTIONS["internal"]))


# ─────────────────────────────────────────
# 共用
# ─────────────────────────────────────────

def summarize_findings(findings: list[dict]) -> dict[str, int]:
    summary = {"不合理": 0, "待人工": 0, "提醒": 0, "錯誤": 0, "其他": 0}
    for f in findings or []:
        sev = f.get("嚴重度", "")
        if sev in summary:
            summary[sev] += 1
        else:
            summary["其他"] += 1
    return summary


def filter_findings(findings: list[dict], opts: dict) -> list[dict]:
    """依 options 過濾要含的 findings。

    一律含: 不合理 + 待人工
    選用: 提醒 (include_reminder), 系統錯誤 (include_system_error)
    """
    keep_sev = {"不合理", "待人工", "提醒"}  # 「提醒」永遠含 (筆誤/錯別字)
    if opts.get("include_reminder"):
        keep_sev.add("其他")  # 額外 catch-all
    if opts.get("include_system_error"):
        keep_sev.add("錯誤")
    return [f for f in (findings or []) if f.get("嚴重度") in keep_sev]


def collect_design_metrics(app_data: dict) -> list[dict]:
    """跑 step3h 算各單元的 HRT/SOR/G 值。"""
    out = []
    try:
        import step3h_design_metrics as dm
    except Exception:
        return out
    for code, unit in sorted((app_data.get("units") or {}).items()):
        try:
            metrics = dm.compute_all_metrics(unit) or {}
        except Exception:
            metrics = {}
        if metrics:
            out.append({
                "代號": code,
                "名稱": unit.get("name_in_doc", ""),
                "標準槽體": unit.get("std_tank", ""),
                **{k: (v.get("value") if isinstance(v, dict) else v) for k, v in metrics.items()},
            })
    return out


def collect_removal_rates(app_data: dict) -> dict[str, dict]:
    """各單元的水質項目削減率 (基於 Σ進質量 vs Σ出質量)。"""
    out = {}
    for code, unit in (app_data.get("units") or {}).items():
        infl = unit.get("influent") or {}
        effl = unit.get("effluent") or {}
        items_all = set()
        for s in list(infl.values()) + list(effl.values()):
            if isinstance(s, dict):
                items_all.update(s.keys())
        items_data = {}
        for item in items_all:
            in_mass = 0.0
            out_mass = 0.0
            in_n = out_n = 0
            for s in infl.values():
                if isinstance(s, dict) and isinstance(s.get(item), dict):
                    try:
                        in_mass += float(s[item].get("質量") or 0)
                        in_n += 1
                    except Exception:
                        pass
            for s in effl.values():
                if isinstance(s, dict) and isinstance(s.get(item), dict):
                    try:
                        out_mass += float(s[item].get("質量") or 0)
                        out_n += 1
                    except Exception:
                        pass
            if in_mass > 0 and out_n > 0:
                items_data[item] = {
                    "進流總質量": round(in_mass, 4),
                    "出流總質量": round(out_mass, 4),
                    "削減率%": round((in_mass - out_mass) / in_mass * 100, 1),
                }
        if items_data:
            out[code] = items_data
    return out


def collect_rule_matches(findings: list[dict]) -> dict[str, int]:
    """各規則被觸發的次數。"""
    from collections import Counter
    ctr = Counter()
    for f in findings or []:
        ctr[f.get("依據", "(無依據)")] += 1
    return dict(ctr.most_common())


def collect_pdf_notes(app_data: dict) -> list[dict]:
    return app_data.get("reviewer_notes") or []


def collect_discharge_standard(business_type: str) -> dict:
    if not business_type:
        return {}
    try:
        from discharge_standards import get_standard
        return get_standard(business_type) or {}
    except Exception:
        return {}


def collect_default_values(business_type: str) -> dict:
    out = {}
    try:
        from step3c_unit_db import DEFAULT_RAW_CONCENTRATIONS, UNIT_DEFAULT_REMOVAL
        out["raw_concentrations"] = DEFAULT_RAW_CONCENTRATIONS
        out["unit_default_removal"] = UNIT_DEFAULT_REMOVAL
    except Exception:
        pass
    return out


# ─────────────────────────────────────────
# 1. Excel (內部覆核)
# ─────────────────────────────────────────

def build_internal_excel(app_data: dict, findings: list[dict] | None = None,
                         options: dict | None = None) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    opts = dict(DEFAULT_OPTIONS["internal"])
    if options:
        opts.update(options)
    findings = filter_findings(findings, opts)
    summary = summarize_findings(findings)

    HEAD_FILL = PatternFill("solid", fgColor="305496")
    HEAD_FONT = Font(bold=True, color="FFFFFF", size=11)
    RED_FILL = PatternFill("solid", fgColor="FFE6E6")
    YEL_FILL = PatternFill("solid", fgColor="FFF2CC")
    GREY_FILL = PatternFill("solid", fgColor="EAEAEA")
    WRAP = Alignment(wrap_text=True, vertical="top", horizontal="left")

    def _head(ws, headers):
        for i, h in enumerate(headers, 1):
            c = ws.cell(1, i, h)
            c.font = HEAD_FONT
            c.fill = HEAD_FILL
            c.alignment = Alignment(horizontal="center")
        ws.freeze_panes = "A2"

    def _autosize(ws, mx=60):
        for col in ws.columns:
            try:
                w = max(len(str(c.value or "")) for c in col)
            except Exception:
                w = 12
            ws.column_dimensions[col[0].column_letter].width = min(max(w + 2, 8), mx)

    wb = openpyxl.Workbook()

    # 1. 摘要
    ws = wb.active
    ws.title = "1.審查摘要"
    rows = [
        ["項目", "值"],
        ["案件名稱", app_data.get("source_pdf", "")],
        ["業別", opts.get("business_type", "(未指定)")],
        ["匯出時間", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
        ["", ""],
        ["━━━ 審查結果統計 ━━━", ""],
        ["🔴 不合理", summary["不合理"]],
        ["🟡 待人工", summary["待人工"]],
        ["💡 提醒/筆誤", summary["提醒"]],
        ["💡 提醒 / 其他", summary["提醒"] + summary["其他"]],
        ["⚠️ 系統錯誤", summary["錯誤"]],
        ["處理單元數", len(app_data.get("units") or {})],
    ]
    for r_idx, r in enumerate(rows, 1):
        for c_idx, v in enumerate(r, 1):
            cell = ws.cell(r_idx, c_idx, v)
            if r_idx == 1:
                cell.font = HEAD_FONT
                cell.fill = HEAD_FILL
    ws.column_dimensions["A"].width = 36
    ws.column_dimensions["B"].width = 50

    # 2. 不合理
    ws = wb.create_sheet("2.不合理🔴")
    _head(ws, ["嚴重度", "類型", "單元", "標準槽體", "對照項目", "描述", "依據"])
    r = 2
    for f in findings:
        if f.get("嚴重度") != "不合理":
            continue
        for i, k in enumerate(["嚴重度", "類型", "單元", "標準槽體", "對照項目", "描述", "依據"], 1):
            cell = ws.cell(r, i, f.get(k, ""))
            cell.alignment = WRAP
            cell.fill = RED_FILL
        r += 1
    _autosize(ws)

    # 3. 待人工
    ws = wb.create_sheet("3.待人工🟡")
    _head(ws, ["嚴重度", "類型", "單元", "標準槽體", "對照項目", "描述", "依據"])
    r = 2
    for f in findings:
        if f.get("嚴重度") != "待人工":
            continue
        for i, k in enumerate(["嚴重度", "類型", "單元", "標準槽體", "對照項目", "描述", "依據"], 1):
            cell = ws.cell(r, i, f.get(k, ""))
            cell.alignment = WRAP
            cell.fill = YEL_FILL
        r += 1
    _autosize(ws)

    # 提醒 / 系統錯誤 (選用)
    other = [f for f in findings if f.get("嚴重度") not in ("不合理", "待人工")]
    if other:
        ws = wb.create_sheet("3b.其他")
        _head(ws, ["嚴重度", "類型", "單元", "標準槽體", "對照項目", "描述", "依據"])
        r = 2
        for f in other:
            for i, k in enumerate(["嚴重度", "類型", "單元", "標準槽體", "對照項目", "描述", "依據"], 1):
                cell = ws.cell(r, i, f.get(k, ""))
                cell.alignment = WRAP
                cell.fill = GREY_FILL
            r += 1
        _autosize(ws)

    # 4. 全廠單元表
    if opts.get("include_unit_table"):
        ws = wb.create_sheet("4.單元表")
        _head(ws, ["代號", "原始名稱", "標準槽體", "進流數", "出流數",
                   "有效容量(m³)", "HRT(hr)", "SOR(m³/m²·d)", "G(1/s)"])
        r = 2
        for code, unit in sorted((app_data.get("units") or {}).items()):
            size = unit.get("size") or {}
            dp = unit.get("design_params") or {}
            def _dp(*keys):
                for k in keys:
                    v = dp.get(k) or {}
                    if isinstance(v, dict):
                        raw = v.get("raw") or v.get("min") or v.get("max")
                        if raw:
                            return str(raw)
                return ""
            ws.cell(r, 1, code)
            ws.cell(r, 2, unit.get("name_in_doc", ""))
            ws.cell(r, 3, unit.get("std_tank", ""))
            ws.cell(r, 4, len(unit.get("influent") or {}))
            ws.cell(r, 5, len(unit.get("effluent") or {}))
            ws.cell(r, 6, size.get("有效容量", ""))
            ws.cell(r, 7, _dp("水力停留時間", "HRT"))
            ws.cell(r, 8, _dp("表面溢流率", "SOR"))
            ws.cell(r, 9, _dp("G值", "速度梯度", "G"))
            r += 1
        _autosize(ws)

    # 5. 各單元水質
    if opts.get("include_water_quality"):
        ws = wb.create_sheet("5.各單元水質")
        _head(ws, ["單元", "方向", "Stream編號", "水質項目", "濃度", "質量"])
        r = 2
        for code, unit in sorted((app_data.get("units") or {}).items()):
            for direction, key in [("進", "influent"), ("出", "effluent")]:
                for stream_code, items in (unit.get(key) or {}).items():
                    if not isinstance(items, dict):
                        continue
                    for item_name, val in items.items():
                        if not isinstance(val, dict):
                            continue
                        ws.cell(r, 1, code)
                        ws.cell(r, 2, direction)
                        ws.cell(r, 3, stream_code)
                        ws.cell(r, 4, item_name)
                        ws.cell(r, 5, val.get("濃度"))
                        ws.cell(r, 6, val.get("質量"))
                        r += 1
        _autosize(ws)

    # 6. 設計參數體檢
    if opts.get("include_design_metrics"):
        dm_rows = collect_design_metrics(app_data)
        if dm_rows:
            ws = wb.create_sheet("6.設計體檢")
            all_keys = ["代號", "名稱", "標準槽體"]
            for row in dm_rows:
                for k in row:
                    if k not in all_keys:
                        all_keys.append(k)
            _head(ws, all_keys)
            r = 2
            for row in dm_rows:
                for i, k in enumerate(all_keys, 1):
                    ws.cell(r, i, row.get(k, ""))
                r += 1
            _autosize(ws)

    # 7. 削減率分項
    if opts.get("include_removal_rate"):
        rr = collect_removal_rates(app_data)
        if rr:
            ws = wb.create_sheet("7.削減率")
            _head(ws, ["單元", "水質項目", "進流總質量", "出流總質量", "削減率%"])
            r = 2
            for code, items in sorted(rr.items()):
                for item, vals in items.items():
                    ws.cell(r, 1, code)
                    ws.cell(r, 2, item)
                    ws.cell(r, 3, vals["進流總質量"])
                    ws.cell(r, 4, vals["出流總質量"])
                    ws.cell(r, 5, vals["削減率%"])
                    r += 1
            _autosize(ws)

    # 8. 觸發規則對照
    if opts.get("include_rule_match"):
        rm = collect_rule_matches(findings)
        if rm:
            ws = wb.create_sheet("8.規則對照")
            _head(ws, ["規則依據", "觸發次數"])
            r = 2
            for rule, cnt in rm.items():
                ws.cell(r, 1, rule).alignment = WRAP
                ws.cell(r, 2, cnt)
                r += 1
            ws.column_dimensions["A"].width = 80
            ws.column_dimensions["B"].width = 12

    # 9. 規則條文節錄
    if opts.get("include_rule_text"):
        try:
            import tank_chemistry as tc
            tc.clear_cache()
            rules = tc.load_rules()
            ws = wb.create_sheet("9.規則庫")
            _head(ws, ["標準槽體", "類型", "加藥類型", "應變動項目", "不應變動項目",
                       "容忍度%", "嚴重度", "學理說明"])
            r = 2
            for name, rule in sorted(rules.items()):
                ws.cell(r, 1, name)
                ws.cell(r, 2, rule.get("類型", ""))
                ws.cell(r, 3, rule.get("加藥類型", ""))
                ws.cell(r, 4, rule.get("應變動原始", ""))
                ws.cell(r, 5, rule.get("不應變動原始", ""))
                ws.cell(r, 6, rule.get("容忍度", ""))
                ws.cell(r, 7, rule.get("嚴重度", ""))
                ws.cell(r, 8, rule.get("學理說明", "")).alignment = WRAP
                r += 1
            _autosize(ws, mx=80)
        except Exception:
            pass

    # 10. PDF 註解
    if opts.get("include_pdf_notes"):
        notes = collect_pdf_notes(app_data)
        if notes:
            ws = wb.create_sheet("10.PDF註解")
            _head(ws, ["頁碼", "作者", "類型", "內容"])
            r = 2
            for n in notes:
                ws.cell(r, 1, n.get("page", ""))
                ws.cell(r, 2, n.get("author", ""))
                ws.cell(r, 3, n.get("subtype", ""))
                ws.cell(r, 4, n.get("contents", "")).alignment = WRAP
                r += 1
            _autosize(ws, mx=80)

    # 11. 業別預設值
    if opts.get("include_default_values"):
        dv = collect_default_values(opts.get("business_type"))
        if dv.get("raw_concentrations"):
            ws = wb.create_sheet("11a.原廢水典型值")
            _head(ws, ["水質項目", "典型值", "單位", "說明"])
            r = 2
            for item, info in dv["raw_concentrations"].items():
                if isinstance(info, dict):
                    ws.cell(r, 1, item)
                    ws.cell(r, 2, info.get("value", ""))
                    ws.cell(r, 3, info.get("unit", ""))
                    ws.cell(r, 4, str(info.get("is_range", "")))
                    r += 1
            _autosize(ws)

    # 12. 放流水標準
    if opts.get("include_discharge_std"):
        ds = collect_discharge_standard(opts.get("business_type"))
        if ds:
            ws = wb.create_sheet("12.放流水標準")
            _head(ws, ["水質項目", "限值", "單位", "備註"])
            r = 2
            for item, val in ds.items():
                if isinstance(val, dict):
                    ws.cell(r, 1, item)
                    ws.cell(r, 2, val.get("value", ""))
                    ws.cell(r, 3, val.get("unit", ""))
                    ws.cell(r, 4, val.get("note", ""))
                else:
                    ws.cell(r, 1, item)
                    ws.cell(r, 2, str(val))
                r += 1
            _autosize(ws)

    # 13. 原始 JSON dump
    if opts.get("include_raw_json_dump"):
        ws = wb.create_sheet("13.JSON原始")
        try:
            ws.cell(1, 1, json.dumps(app_data, ensure_ascii=False, indent=2)[:32760])
            ws.column_dimensions["A"].width = 120
            ws.cell(1, 1).alignment = WRAP
        except Exception:
            ws.cell(1, 1, "(JSON 太大或序列化失敗)")

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─────────────────────────────────────────
# 2. Word (廠商通知)
# ─────────────────────────────────────────

def build_vendor_word(app_data: dict, findings: list[dict] | None = None,
                      options: dict | None = None) -> bytes:
    try:
        from docx import Document
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("未安裝 python-docx。請執行: pip install python-docx 後再試。")

    opts = dict(DEFAULT_OPTIONS["vendor"])
    if options:
        opts.update(options)
    findings = filter_findings(findings, opts)

    doc = Document()
    title = doc.add_heading("水污染防治措施申請書 - 審查意見書", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("案件名稱: ").bold = True
    p.add_run(str(app_data.get("source_pdf", "")) + "\n")
    p.add_run("業別: ").bold = True
    p.add_run(str(opts.get("business_type", "(未指定)")) + "\n")
    p.add_run("審查日期: ").bold = True
    p.add_run(datetime.now().strftime("%Y-%m-%d"))

    # 一、待修正項目
    red = [f for f in findings if f.get("嚴重度") == "不合理"]
    doc.add_heading(f"一、待修正項目 ({len(red)} 項)", level=1)
    if not red:
        doc.add_paragraph("(本案無此類項目)")
    for f in red:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(f"{f.get('單元', '')} / {f.get('對照項目', '')}")
        r.bold = True
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        doc.add_paragraph(f"問題: {f.get('描述', '')}")
        doc.add_paragraph(f"依據: {f.get('依據', '')}")
        doc.add_paragraph("廠商回覆: " + "_" * 60)

    # 二、待澄清項目
    yel = [f for f in findings if f.get("嚴重度") == "待人工"]
    doc.add_heading(f"二、待澄清項目 ({len(yel)} 項)", level=1)
    if not yel:
        doc.add_paragraph("(本案無此類項目)")
    for f in yel:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(f"{f.get('單元', '')} / {f.get('對照項目', '')}")
        r.bold = True
        r.font.color.rgb = RGBColor(0xBF, 0x90, 0x00)
        doc.add_paragraph(f"說明: {f.get('描述', '')}")
        doc.add_paragraph("廠商說明: " + "_" * 60)

    # 三、單元表
    if opts.get("include_unit_table"):
        doc.add_heading("三、處理單元一覽", level=1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Light Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "代號"
        hdr[1].text = "名稱"
        hdr[2].text = "標準槽體"
        hdr[3].text = "有效容量(m³)"
        for code, u in sorted((app_data.get("units") or {}).items()):
            row = tbl.add_row().cells
            row[0].text = code
            row[1].text = u.get("name_in_doc", "")
            row[2].text = u.get("std_tank", "")
            row[3].text = str((u.get("size") or {}).get("有效容量", ""))

    # 四、規則對照
    if opts.get("include_rule_match"):
        rm = collect_rule_matches(findings)
        if rm:
            doc.add_heading("四、本次觸發規則依據", level=1)
            for rule, cnt in rm.items():
                doc.add_paragraph(f"• {rule} (觸發 {cnt} 次)", style="List Bullet")

    # 五、規則條文節錄
    if opts.get("include_rule_text"):
        doc.add_heading("五、相關規則條文", level=1)
        try:
            import tank_chemistry as tc
            tc.clear_cache()
            rules = tc.load_rules()
            for name, rule in sorted(rules.items()):
                desc = rule.get("學理說明", "")
                if desc:
                    doc.add_paragraph(f"{name}: {desc}", style="List Bullet")
        except Exception:
            doc.add_paragraph("(無法載入規則庫)")

    # 六、放流水標準
    if opts.get("include_discharge_std"):
        ds = collect_discharge_standard(opts.get("business_type"))
        if ds:
            doc.add_heading(f"六、放流水標準 ({opts.get('business_type', '')})", level=1)
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Light Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text = "水質項目"
            hdr[1].text = "限值"
            hdr[2].text = "單位"
            for item, val in ds.items():
                row = tbl.add_row().cells
                row[0].text = str(item)
                if isinstance(val, dict):
                    row[1].text = str(val.get("value", ""))
                    row[2].text = str(val.get("unit", ""))
                else:
                    row[1].text = str(val)

    # 簽核
    doc.add_paragraph()
    doc.add_paragraph("─" * 40)
    p = doc.add_paragraph()
    p.add_run("審查人: __________________  日期: __________\n").bold = True
    p.add_run("覆核人: __________________  日期: __________").bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─────────────────────────────────────────
# 3. JSON (整合)
# ─────────────────────────────────────────

def build_integrated_json(app_data: dict, findings: list[dict] | None = None,
                          options: dict | None = None) -> bytes:
    opts = dict(DEFAULT_OPTIONS["json"])
    if options:
        opts.update(options)
    findings = filter_findings(findings, opts)
    summary = summarize_findings(findings)
    bt = opts.get("business_type", "")

    out = {
        "案件": app_data.get("source_pdf", ""),
        "業別": bt,
        "匯出時間": datetime.now().isoformat(),
        "統計": {
            "不合理": summary["不合理"],
            "待人工": summary["待人工"],
            "提醒": summary["提醒"],
            "錯誤": summary["錯誤"],
            "處理單元數": len(app_data.get("units") or {}),
        },
        "findings": findings,
    }

    if opts.get("include_unit_table"):
        out["units"] = app_data.get("units") or {}
    if opts.get("include_design_metrics"):
        out["design_metrics"] = collect_design_metrics(app_data)
    if opts.get("include_removal_rate"):
        out["removal_rates"] = collect_removal_rates(app_data)
    if opts.get("include_rule_match"):
        out["rule_matches"] = collect_rule_matches(findings)
    if opts.get("include_pdf_notes"):
        out["pdf_notes"] = collect_pdf_notes(app_data)
    if opts.get("include_default_values"):
        out["default_values"] = collect_default_values(bt)
    if opts.get("include_discharge_std"):
        out["discharge_standard"] = collect_discharge_standard(bt)
    if opts.get("include_raw_json_dump"):
        out["_raw_app_data"] = app_data

    return json.dumps(out, ensure_ascii=False, indent=2).encode("utf-8")


# ─────────────────────────────────────────
# 主入口
# ─────────────────────────────────────────

def build_export(target: str, app_data: dict, findings: list[dict] | None = None,
                 options: dict | None = None,
                 base_name: str = "report") -> tuple[bytes, str, str]:
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    if target == "internal":
        data = build_internal_excel(app_data, findings, options)
        fname = f"{base_name}_內部覆核_{ts}.xlsx"
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif target == "vendor":
        data = build_vendor_word(app_data, findings, options)
        fname = f"{base_name}_廠商通知_{ts}.docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif target == "json":
        data = build_integrated_json(app_data, findings, options)
        fname = f"{base_name}_整合資料_{ts}.json"
        mime = "application/json"
    else:
        raise ValueError(f"未知 target: {target}")
    return data, fname, mime


if __name__ == "__main__":
    import os
    import sys
    base = os.path.dirname(os.path.abspath(__file__))
    jsons = sorted([f for f in os.listdir(base) if f.startswith("application_") and f.endswith(".json")])
    if not jsons:
        print("找不到 application_*.json")
        sys.exit(0)
    with open(os.path.join(base, jsons[-1]), encoding="utf-8") as f:
        app = json.load(f)
    from step3d_principle_check import run_advanced_checks
    findings = run_advanced_checks(app, business_type="電鍍業")
    print(f"案件: {app.get('source_pdf')}, findings: {len(findings)}")
    for tgt in ["internal", "json"]:
        try:
            data, fname, mime = build_export(tgt, app, findings,
                                             options={"business_type": "電鍍業"},
                                             base_name="test")
            print(f"  {tgt}: {len(data)/1024:.0f} KB → {fname}")
            with open(os.path.join(base, fname), "wb") as f:
                f.write(data)
        except Exception as e:
            import traceback
            print(f"  {tgt}: ERROR {e}")
            traceback.print_exc()

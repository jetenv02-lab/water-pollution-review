# -*- coding: utf-8 -*-
"""Step 4: 把 step3 的比對結果產出最終輸出。

輸出:
  1) 審查報告/{案名}_比對結果.xlsx
     - 總表分頁(全部 findings 一張表)
     - 每個申請單元一張分頁(該單元觸發的所有規則)
  2) 審查報告/{案名}_審查意見.docx
     - 格式化意見書,可直接交付技師
  3) 寫回 規則庫.xlsx 的 _審查紀錄 分頁

依賴: openpyxl, python-docx (若未安裝會跳過 Word 輸出)
用法: python step4_output.py "comparison_xxx.json"
"""
import json
import os
import sys
from datetime import datetime
from collections import defaultdict
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

BASE = r"C:\Users\jeten\Desktop\AI\水措審查"
RULES_XLSX = os.path.join(BASE, "規則庫.xlsx")
REPORT_DIR = os.path.join(BASE, "審查報告")
os.makedirs(REPORT_DIR, exist_ok=True)

# 樣式
HEADER_FILL = PatternFill("solid", fgColor="2F5496")
HEADER_FONT = Font(color="FFFFFF", bold=True, size=11)
SEVERITY_FILL = {
    "不合理": PatternFill("solid", fgColor="FCE4D6"),  # 淺紅
    "待人工": PatternFill("solid", fgColor="FFF2CC"),  # 淺黃
    "合理": PatternFill("solid", fgColor="E2EFDA"),    # 淺綠
}
WRAP = Alignment(wrap_text=True, vertical="top")
CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)
THIN = Side(style="thin", color="BFBFBF")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def output_excel(comparison, out_path):
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    findings = comparison["findings"]
    fields = ["申請單元", "標準槽體", "缺失ID", "規則來源", "檢查類型",
              "對照項目", "判定", "描述", "規則原文", "原文缺失"]
    widths = [12, 14, 10, 24, 12, 14, 10, 36, 36, 36]

    # 1) 總表
    ws = wb.create_sheet("_總表")
    ws.append(fields)
    for r in findings:
        row = [r.get(f, "") for f in fields]
        ws.append(row)
        # 套色
        fill = SEVERITY_FILL.get(r.get("判定"))
        if fill:
            for c in range(1, len(fields) + 1):
                ws.cell(row=ws.max_row, column=c).fill = fill
    # 樣式
    for c in range(1, len(fields) + 1):
        cell = ws.cell(row=1, column=c)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = CENTER
        cell.border = BORDER
        ws.column_dimensions[get_column_letter(c)].width = widths[c - 1]
    ws.freeze_panes = "A2"
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=len(fields)):
        for cell in row:
            cell.alignment = WRAP
            cell.border = BORDER

    # 2) 每個申請單元一張分頁
    by_unit = defaultdict(list)
    for r in findings:
        by_unit[r["申請單元"]].append(r)
    for unit, rows in sorted(by_unit.items()):
        sname = unit[:31].replace("/", "-")
        ws = wb.create_sheet(sname)
        ws.append(fields)
        for r in rows:
            row = [r.get(f, "") for f in fields]
            ws.append(row)
            fill = SEVERITY_FILL.get(r.get("判定"))
            if fill:
                for c in range(1, len(fields) + 1):
                    ws.cell(row=ws.max_row, column=c).fill = fill
        for c in range(1, len(fields) + 1):
            cell = ws.cell(row=1, column=c)
            cell.fill = HEADER_FILL
            cell.font = HEADER_FONT
            cell.alignment = CENTER
            cell.border = BORDER
            ws.column_dimensions[get_column_letter(c)].width = widths[c - 1]
        ws.freeze_panes = "A2"
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=len(fields)):
            for cell in row:
                cell.alignment = WRAP
                cell.border = BORDER

    wb.save(out_path)
    print(f"已輸出 Excel: {out_path}")


def output_word(comparison, out_path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        print("未安裝 python-docx,跳過 Word 輸出。安裝: pip install python-docx")
        return

    doc = Document()
    doc.add_heading("水措審查意見書", level=0)
    doc.add_paragraph(f"審查文件: {comparison.get('application', '')}")
    doc.add_paragraph(f"審查時間: {comparison.get('compared_at', '')}")
    doc.add_paragraph(f"比對總數: {comparison.get('total_findings', 0)}")

    findings = comparison["findings"]
    not_ok = [r for r in findings if r["判定"] == "不合理"]
    manual = [r for r in findings if r["判定"] == "待人工"]

    doc.add_heading(f"一、自動判定不合理項目 ({len(not_ok)} 項)", level=1)
    if not_ok:
        for i, r in enumerate(not_ok, 1):
            p = doc.add_paragraph(style="List Number")
            p.add_run(f"【{r['申請單元']} {r['標準槽體']}】").bold = True
            p.add_run(f" 對照項目: {r['對照項目']} | ")
            p.add_run(f"描述: {r['描述']}\n")
            p.add_run(f"規則原文: {r['規則原文']}\n").italic = True
            p.add_run(f"出處: {r['規則來源']} (缺失ID {r['缺失ID']})").font.size = Pt(9)
    else:
        doc.add_paragraph("(無)")

    doc.add_heading(f"二、待人工複核項目 ({len(manual)} 項)", level=1)
    by_unit = defaultdict(list)
    for r in manual:
        by_unit[r["申請單元"]].append(r)
    for unit, rows in sorted(by_unit.items()):
        doc.add_heading(f"{unit} ({rows[0]['標準槽體']})", level=2)
        for r in rows:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{r['檢查類型']}/{r['對照項目']}] {r['規則原文']}")

    doc.save(out_path)
    print(f"已輸出 Word: {out_path}")


def update_review_log(comparison, xlsx_paths):
    """寫回 規則庫.xlsx 的 _審查紀錄。"""
    if not os.path.exists(RULES_XLSX):
        return
    wb = openpyxl.load_workbook(RULES_XLSX)
    if "_審查紀錄" not in wb.sheetnames:
        return
    ws = wb["_審查紀錄"]
    not_ok_count = sum(1 for r in comparison["findings"] if r["判定"] == "不合理")
    result_text = "不合格" if not_ok_count > 0 else "合格(自動)"
    # 計算同一份 application 的審查次數
    app = comparison.get("application", "")
    count = 1
    for r in range(2, ws.max_row + 1):
        if str(ws.cell(row=r, column=1).value or "").strip() == app:
            count += 1
    ws.append([
        app, not_ok_count, result_text,
        datetime.now().strftime("%Y-%m-%d %H:%M"),
        count, "; ".join(xlsx_paths)
    ])
    wb.save(RULES_XLSX)
    print(f"已寫回 _審查紀錄: {app} (第 {count} 次審查)")


def main():
    if len(sys.argv) >= 2:
        comp_path = sys.argv[1]
    else:
        comps = sorted([f for f in os.listdir(BASE) if f.startswith("comparison_") and f.endswith(".json")])
        if not comps:
            print("找不到 comparison_*.json,請先跑 step3")
            return
        comp_path = os.path.join(BASE, comps[-1])
    if not os.path.exists(comp_path):
        print(f"找不到: {comp_path}")
        return

    with open(comp_path, "r", encoding="utf-8") as f:
        comparison = json.load(f)

    case = os.path.splitext(os.path.basename(comparison["application"]))[0]
    xlsx_out = os.path.join(REPORT_DIR, f"{case}_比對結果.xlsx")
    docx_out = os.path.join(REPORT_DIR, f"{case}_審查意見.docx")

    output_excel(comparison, xlsx_out)
    output_word(comparison, docx_out)
    update_review_log(comparison, [xlsx_out, docx_out])


if __name__ == "__main__":
    main()
